"""
Document content extractors for PDF, Excel (.xlsx), and Word (.docx) files.

Each extractor returns a normalised `DocumentContent` dataclass so the rest of
the pipeline is format-agnostic.
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field
from typing import List, Optional

import pdfplumber
import openpyxl
from docx import Document as DocxDocument


@dataclass
class TableData:
    sheet_or_page: str
    headers: List[str]
    rows: List[List[str]]


@dataclass
class DocumentContent:
    filename: str
    format: str
    text_blocks: List[str] = field(default_factory=list)
    tables: List[TableData] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


def _clean(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def extract_pdf(path: str) -> DocumentContent:
    doc = DocumentContent(filename=os.path.basename(path), format="pdf")
    with pdfplumber.open(path) as pdf:
        doc.metadata = pdf.metadata or {}
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text:
                doc.text_blocks.append(text)

            for table in page.extract_tables():
                if not table:
                    continue
                headers = [_clean(c) for c in table[0]]
                rows = [[_clean(c) for c in row] for row in table[1:]]
                doc.tables.append(
                    TableData(
                        sheet_or_page=f"Page {page_num}",
                        headers=headers,
                        rows=rows,
                    )
                )
    return doc


# ---------------------------------------------------------------------------
# Excel (.xlsx)
# ---------------------------------------------------------------------------

def extract_excel(path: str) -> DocumentContent:
    doc = DocumentContent(filename=os.path.basename(path), format="xlsx")
    wb = openpyxl.load_workbook(path, data_only=True)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_raw = list(ws.iter_rows(values_only=True))
        if not rows_raw:
            continue

        headers = [_clean(c) for c in rows_raw[0]]
        rows = [[_clean(c) for c in row] for row in rows_raw[1:]]

        all_empty = all(all(cell == "" for cell in row) for row in rows)
        if all_empty and all(h == "" for h in headers):
            continue

        doc.tables.append(
            TableData(sheet_or_page=sheet_name, headers=headers, rows=rows)
        )

    return doc


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------

def extract_word(path: str) -> DocumentContent:
    doc_content = DocumentContent(filename=os.path.basename(path), format="docx")
    docx = DocxDocument(path)

    doc_content.metadata = {
        "author": docx.core_properties.author or "",
        "title": docx.core_properties.title or "",
    }

    for para in docx.paragraphs:
        text = para.text.strip()
        if text:
            doc_content.text_blocks.append(text)

    for table_idx, table in enumerate(docx.tables):
        rows_raw = []
        for row in table.rows:
            rows_raw.append([_clean(cell.text) for cell in row.cells])
        if not rows_raw:
            continue
        headers = rows_raw[0]
        rows = rows_raw[1:]
        doc_content.tables.append(
            TableData(
                sheet_or_page=f"Table {table_idx + 1}",
                headers=headers,
                rows=rows,
            )
        )

    return doc_content


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {
    ".pdf": extract_pdf,
    ".xlsx": extract_excel,
    ".docx": extract_word,
}


def extract(path: str) -> DocumentContent:
    ext = os.path.splitext(path)[1].lower()
    extractor = SUPPORTED_EXTENSIONS.get(ext)
    if extractor is None:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )
    return extractor(path)
